import os
import re
import threading
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml, OxmlElement

# ==========================================
# 后端：全功能学术排版引擎 V27.0 (三线表专项修复)
# ==========================================
class WordFormatterEngine:
    def __init__(self, input_path, output_path, config, media_config, page_config):
        self.input_path = input_path
        self.output_path = output_path
        self.config = config
        self.media_config = media_config
        self.page_config = page_config
        self.align_map = {"左对齐": 0, "居中": 1, "右对齐": 2, "两端对齐": 3}

    def is_h1(self, text):
        strict_p = r'^\s*(第[一二三四五六七八九十百零]+章|摘要|Abstract|致谢|参考文献|绪论|引言|附录)'
        if bool(re.match(strict_p, text, re.I)): return True
        if len(text) < 40 and any(kw in text for kw in ["发表", "成果", "学术论文"]): return True
        return False

    def is_body_start(self, text):
        return bool(re.match(r'^\s*(第[一二三四五六七八九十百零]+章|绪论|引言)', text))

    def has_toc(self, doc):
        for p in doc.paragraphs:
            if 'TOC' in p._element.xml: return True
        return False

    def insert_toc_field(self, paragraph):
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin'); run._r.append(fldChar1)
        instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = ' TOC \\o "1-3" \\h \\z \\u '; run._r.append(instrText)
        fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate'); run._r.append(fldChar2)
        run.add_text("【目录指令已生成，请在Word中右键更新域】")
        fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end'); run._r.append(fldChar3)

    def set_section_paging(self, section, fmt, start_val=None):
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn('w:pgNumType'))
        if pgNumType is None:
            pgNumType = OxmlElement('w:pgNumType')
            sectPr.append(pgNumType)
        pgNumType.set(qn('w:fmt'), fmt)
        if start_val is not None: pgNumType.set(qn('w:start'), str(start_val))

    # --- 核心：三线表暴力重画逻辑 V27.0 ---
    def process_table(self, table):
        if not self.media_config.get('three_line'): return
        
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        # 1. 彻底切断表格样式联系
        style = tblPr.xpath('./w:tblStyle')
        if style: style[0].set(qn('w:val'), 'TableNormal')
        
        # 2. 暴力清除表格级别所有边框 (防止残留)
        old_tbl_borders = tblPr.xpath('./w:tblBorders')
        if old_tbl_borders: tblPr.remove(old_tbl_borders[0])
        
        new_tbl_borders = OxmlElement('w:tblBorders')
        for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            elm = OxmlElement(f'w:{b}')
            elm.set(qn('w:val'), 'nil') # 全体设为空
            new_tbl_borders.append(elm)
        tblPr.append(new_tbl_borders)

        # 3. 逐个单元格精准画线
        rows = table.rows
        num_rows = len(rows)
        for r_idx, row in enumerate(rows):
            for cell in row.cells:
                # 垂直/段落居中
                cell.vertical_alignment = 1
                for p in cell.paragraphs: p.alignment = 1
                
                tcPr = cell._element.get_or_add_tcPr()
                # 清除单元格边框残留
                old_tc_borders = tcPr.xpath('./w:tcBorders')
                if old_tc_borders: tcPr.remove(old_tc_borders[0])
                
                tcB = OxmlElement('w:tcBorders')
                for side in ['top', 'bottom', 'left', 'right']:
                    b_elm = OxmlElement(f'w:{side}')
                    b_elm.set(qn('w:val'), 'nil') # 默认无
                    
                    # 顶行顶线: 1.5pt (12)
                    if r_idx == 0 and side == 'top':
                        b_elm.set(qn('w:val'), 'single'); b_elm.set(qn('w:sz'), '12'); b_elm.set(qn('w:color'), 'auto')
                    # 顶行底线 (表头线): 0.75pt (6)
                    if r_idx == 0 and side == 'bottom':
                        b_elm.set(qn('w:val'), 'single'); b_elm.set(qn('w:sz'), '6'); b_elm.set(qn('w:color'), 'auto')
                    # 末行底线: 1.5pt (12)
                    if r_idx == num_rows - 1 and side == 'bottom':
                        b_elm.set(qn('w:val'), 'single'); b_elm.set(qn('w:sz'), '12'); b_elm.set(qn('w:color'), 'auto')
                    
                    tcB.append(b_elm)
                tcPr.append(tcB)

    def run(self):
        doc = Document(self.input_path)
        
        # 1. 物理分节
        if self.page_config.get('enabled'):
            if self.page_config.get('auto_toc') and not self.has_toc(doc):
                for p in doc.paragraphs:
                    if self.is_h1(p.text):
                        target = p.insert_paragraph_before()
                        target.text = "目录"; target.alignment = 1
                        self.apply_style(target, self.config['h1'])
                        self.insert_toc_field(doc.add_paragraph())
                        break
            for para in doc.paragraphs:
                if self.is_body_start(para.text):
                    new_p = para.insert_paragraph_before()
                    sectPr = OxmlElement('w:sectPr')
                    type_obj = OxmlElement('w:type'); type_obj.set(qn('w:val'), 'nextPage')
                    sectPr.append(type_obj); new_p._element.get_or_add_pPr().append(sectPr)
                    break

        # 2. 标题/公式/样式
        chap_cnt, eq_cnt, state = 0, 0, "body"
        for para in doc.paragraphs:
            txt = para.text.strip()
            if not txt and not para._element.xpath('.//m:oMath'): continue
            if self.is_h1(txt):
                chap_cnt += 1; eq_cnt = 0
                state = "abstract" if "摘要" in txt or "Abstract" in txt else ("reference" if "参考" in txt or "成果" in txt else "body")
                if self.media_config.get('auto_pb') and chap_cnt > 1:
                    para.insert_paragraph_before().add_run().add_break(WD_BREAK.PAGE)
                self.apply_style(para, self.config['h1'])
                continue
            if bool(re.match(r'^\s*\d+\.\d+\.\d+', txt)):
                self.apply_style(para, self.config['h3']); state = "body"; continue
            if bool(re.match(r'^\s*\d+\.\d+', txt)):
                self.apply_style(para, self.config['h2']); state = "body"; continue
            if re.match(r'^\s*(关键词|Keywords?)[:：]', txt, re.I):
                self.apply_style(para, self.config['keyword'])
            elif bool(re.match(r'^\s*(图|表|Fig|Table)[\s\d]', txt, re.I)):
                self.apply_style(para, self.config['caption'])
            elif self.media_config.get('eq_enable') and (para._element.xpath('.//m:oMath') or (len(txt)<60 and '=' in txt)):
                eq_cnt += 1; para.alignment = 1; para.paragraph_format.first_line_indent = 0
                para.add_run(f"\t#( {chap_cnt}-{eq_cnt} )")
                para.paragraph_format.tab_stops.add_tab_stop(Cm(8), 1); para.paragraph_format.tab_stops.add_tab_stop(Cm(16), 2)
            else:
                target = self.config['abstract'] if state == "abstract" else (self.config['ref_body'] if state == "reference" else self.config['body'])
                self.apply_style(para, target)

        # 3. 媒体 (保持居中修复)
        w_s, w_d = float(self.media_config.get('img_s', 8.0)), float(self.media_config.get('img_d', 14.0))
        for s in doc.inline_shapes: s.width = Cm(w_d) if s.width.cm > 10.0 else Cm(w_s)
        if self.media_config.get('center_all'):
            for p in doc.paragraphs:
                if p._element.xpath('.//w:drawing') or p._element.xpath('.//w:pict'): p.alignment = 1

        # 4. 表格转换 (调用最新修复版)
        for t in doc.tables:
            self.process_table(t)

        # 5. 页眉页脚
        if self.page_config.get('enabled'):
            for i, section in enumerate(doc.sections):
                hp = section.header.paragraphs[0]; hp.text = self.page_config.get('header_text', ""); hp.alignment = 1
                if self.page_config.get('header_line'):
                    pPr = hp._element.get_or_add_pPr()
                    pBdr = OxmlElement('w:pBdr')
                    bottom = OxmlElement('w:bottom'); bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6'); bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'auto')
                    pBdr.append(bottom); pPr.append(pBdr)
                if self.page_config.get('show_page_num'):
                    fp = section.footer.paragraphs[0]; fp.text = ""; fp.alignment = 1
                    run = fp.add_run()
                    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin'); run._r.append(f1)
                    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = "PAGE"; run._r.append(it)
                    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end'); run._r.append(f2)
                fmt = 'decimal' if self.page_config.get('toc_no_roman') else 'romanLower'
                if i == 0: self.set_section_paging(section, fmt)
                else: self.set_section_paging(section, 'decimal', start_val=1)

        doc.save(self.output_path)

    def apply_style(self, para, conf):
        if not conf.get('enabled', True): return
        para.alignment = self.align_map.get(conf.get('align', "居中"), 1)
        pf = para.paragraph_format
        fs = float(conf.get('font_size', 12))
        indent = float(conf.get('first_indent', 0))
        pf.first_line_indent = Pt(fs * indent) if indent > 0 else 0
        pf.line_spacing = 1.5
        for r in para.runs:
            r.font.size = Pt(fs); r.bold = conf.get('bold', False)
            r.font.name = conf.get('en_font', 'Times New Roman'); r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), conf['zh_font'])

# ==========================================
# 前端：GUI 界面
# ==========================================
class App:
    def __init__(self, root):
        self.root = root
        self.root.title("Word 学术排版终极版 V27.0")
        self.root.geometry("980x920")
        self.style_vars = {}
        self.setup_ui()

    def setup_ui(self):
        f_frame = ttk.LabelFrame(self.root, text="文件设置", padding=10)
        f_frame.pack(fill=tk.X, padx=10, pady=5)
        self.path_var = tk.StringVar()
        ttk.Entry(f_frame, textvariable=self.path_var, width=90).pack(side=tk.LEFT, padx=5)
        ttk.Button(f_frame, text="选择文件", command=self.select_file).pack(side=tk.LEFT)

        nb = ttk.Notebook(self.root)
        nb.pack(fill=tk.BOTH, expand=True, padx=10)

        sections = [
            ('h1','一级标题', "黑体", "Arial", 16, True, "居中", 0),
            ('h2','二级标题', "黑体", "Arial", 14, True, "左对齐", 0),
            ('h3','三级标题', "黑体", "Arial", 12, True, "左对齐", 0),
            ('abstract','摘要内容', "楷体", "Times New Roman", 12, False, "两端对齐", 2),
            ('keyword','关键词行', "宋体", "Times New Roman", 12, False, "左对齐", 0),
            ('body','普通正文', "宋体", "Times New Roman", 12, False, "两端对齐", 2),
            ('ref_body','参考/成果正文', "宋体", "Times New Roman", 10.5, False, "左对齐", 0),
            ('caption','图注表题', "楷体", "Times New Roman", 11, False, "居中", 0) 
        ]
        for k, name, zh, en, size, bold, align, indent in sections:
            f = ttk.Frame(nb, padding=15); nb.add(f, text=name)
            self.style_vars[k] = self.create_form(f, zh, en, size, bold, align, indent)

        self.adv_tab = ttk.Frame(nb, padding=15); nb.add(self.adv_tab, text="公式与媒体")
        self.adv_vars = self.create_adv_form(self.adv_tab)

        self.page_tab = ttk.Frame(nb, padding=15); nb.add(self.page_tab, text="页眉页脚与页码")
        self.page_vars = self.create_page_form(self.page_tab)

        ttk.Button(self.root, text="🚀 启动 V27.0 强力排版 (表格修复版)", command=self.run_thread).pack(pady=10)

    def create_form(self, master, zh, en, size, bold, align, indent):
        v = {'enabled': tk.BooleanVar(value=True)}
        ttk.Checkbutton(master, text="启用调整", variable=v['enabled']).grid(row=0, column=0, sticky="w", pady=5)
        ttk.Label(master, text="中文字体:").grid(row=1, column=0); v['zh_font'] = ttk.Combobox(master, values=["宋体", "黑体", "楷体"], width=12); v['zh_font'].set(zh); v['zh_font'].grid(row=1, column=1)
        ttk.Label(master, text="西文字体:").grid(row=1, column=2); v['en_font'] = ttk.Combobox(master, values=["Times New Roman", "Arial"], width=12); v['en_font'].set(en); v['en_font'].grid(row=1, column=3)
        ttk.Label(master, text="字号:").grid(row=2, column=0); v['font_size'] = ttk.Spinbox(master, from_=5, to=72, width=11); v['font_size'].set(size); v['font_size'].grid(row=2, column=1)
        ttk.Label(master, text="对齐:").grid(row=2, column=2); v['align'] = ttk.Combobox(master, values=["左对齐", "居中", "两端对齐"], width=12); v['align'].set(align); v['align'].grid(row=2, column=3)
        ttk.Label(master, text="缩进:").grid(row=3, column=0); v['first_indent'] = ttk.Spinbox(master, from_=0, to=8, increment=1, width=11); v['first_indent'].set(indent); v['first_indent'].grid(row=3, column=1)
        v['bold'] = tk.BooleanVar(value=bold); ttk.Checkbutton(master, text="加粗", variable=v['bold']).grid(row=3, column=2)
        return v

    def create_adv_form(self, master):
        v = {'eq_enable': tk.BooleanVar(value=True), 'img_s': tk.DoubleVar(value=8.0), 'img_d': tk.DoubleVar(value=14.0), 'auto_pb': tk.BooleanVar(value=True), 'three_line': tk.BooleanVar(value=True), 'center_all': tk.BooleanVar(value=True)}
        ttk.Checkbutton(master, text="开启公式识别与对齐", variable=v['eq_enable']).grid(row=0, column=0, columnspan=2, sticky="w", pady=5)
        ttk.Label(master, text="单栏图宽:").grid(row=1, column=0); ttk.Spinbox(master, from_=0, to=15, textvariable=v['img_s'], width=10).grid(row=1, column=1)
        ttk.Label(master, text="双栏图宽:").grid(row=1, column=2); ttk.Spinbox(master, from_=0, to=25, textvariable=v['img_d'], width=10).grid(row=1, column=3)
        ttk.Checkbutton(master, text="强制居中图片/表格", variable=v['center_all']).grid(row=2, column=0, pady=5, sticky="w")
        ttk.Checkbutton(master, text="H1强制分页", variable=v['auto_pb']).grid(row=2, column=2, pady=5, sticky="w")
        ttk.Checkbutton(master, text="标准三线表转换", variable=v['three_line']).grid(row=3, column=0, pady=5, sticky="w")
        return v

    def create_page_form(self, master):
        v = {'enabled': tk.BooleanVar(value=True), 'auto_toc': tk.BooleanVar(value=True), 'toc_no_roman': tk.BooleanVar(value=False), 'header_text': tk.StringVar(value="毕业论文排版测试"), 'header_line': tk.BooleanVar(value=True), 'show_page_num': tk.BooleanVar(value=True)}
        ttk.Checkbutton(master, text="分节页码总控", variable=v['enabled']).grid(row=0, column=0, columnspan=2, sticky="w", pady=5)
        ttk.Checkbutton(master, text="自动生成目录引导", variable=v['auto_toc']).grid(row=1, column=0, sticky="w")
        ttk.Checkbutton(master, text="前置部分不使用罗马页码", variable=v['toc_no_roman']).grid(row=1, column=1, sticky="w")
        ttk.Label(master, text="页眉内容:").grid(row=2, column=0); ttk.Entry(master, textvariable=v['header_text'], width=40).grid(row=2, column=1, pady=5)
        ttk.Checkbutton(master, text="显示页眉下横线", variable=v['header_line']).grid(row=3, column=0, sticky="w")
        ttk.Checkbutton(master, text="显示页脚居中页码", variable=v['show_page_num']).grid(row=3, column=1, sticky="w")
        return v

    def select_file(self):
        p = filedialog.askopenfilename(filetypes=[("Word", "*.docx")]); 
        if p: self.path_var.set(p)

    def run_thread(self):
        p = self.path_var.get()
        if not p: return messagebox.showerror("错误", "请选择文件")
        def task():
            try:
                cfg = {k: {sk: (sv.get() if hasattr(sv, 'get') else sv) for sk, sv in d.items()} for k, d in self.style_vars.items()}
                adv = {k: v.get() for k, v in self.adv_vars.items()}
                page = {k: v.get() for k, v in self.page_vars.items()}
                out = os.path.join(os.path.dirname(p), f"表格修复版_{os.path.basename(p)}")
                WordFormatterEngine(p, out, cfg, adv, page).run()
                messagebox.showinfo("成功", f"排版完成！表格已重新机械绘制。\n\n新文件：{out}")
            except Exception as e: messagebox.showerror("异常", f"详情：{str(e)}")
        threading.Thread(target=task, daemon=True).start()

if __name__ == "__main__":
    root = tk.Tk(); App(root); root.mainloop()
